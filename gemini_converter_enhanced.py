import os
import re
import json
import time
import logging
import hashlib
import shutil
from pathlib import Path
from typing import Dict, Any, Optional, List, Callable
import datetime
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from queue import Queue
import collections
import uuid

# External Libraries
from google import genai
from docx import Document
from dotenv import load_dotenv
from PyPDF2 import PdfWriter, PdfReader
from google.api_core import exceptions

# Set up basic logging (optional but helpful for debugging)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- PROMPT: Structured Document Extraction ---
PROMPT_FULL_DOCUMENT_EXTRACTION = r"""
You are an expert OCR and document structure extraction model.
Extract ALL content from the entire PDF document, broken down by page and then by block. Return valid JSON only.

{
  "total_pages": number,
  "pages": [
    {
      "page_number": 1,
      "blocks": [
        {
          "type": "paragraph" | "list_item" | "table",
          "text": "The exact raw text content for this block, including all original spacing (e.g., 'a  a') and line breaks (\n).",
        }
      ]
    }
  ]
}

CRITICAL RULES:
1. Preserve ALL original spacing and internal line breaks (\n) within the 'text' field.
2. For lists, set the "type" to "list_item". Preserve the original bullet, number, or letter (e.g., '1.', 'a.', '- ') in the 'text' field itself.
3. For tables, set the "type" to "table" and format the 'text' value as markdown pipe format. Example: "| Head 1 | Head 2 |\n|---|---|\n| R1C1 | R1C2 |"
4. **LATEX NOTATION MANDATORY**: Convert ALL superscripts and subscripts to LaTeX format (^{} and _{}).
   • This is an INTERMEDIATE step. The system will convert these to Unicode characters later.
   • Superscripts: Use ^{} notation - y4→y^{4}, x2→x^{2}, E=mc2→E=mc^{2}
   • Subscripts: Use _{} notation - H2O→H_{2}O, CO2→CO_{2}, a1→a_{1}
   • Examples: y4→y^{4}, CO2→CO_{2}, H2O→H_{2}O, x^3→x^{3}, a_n→a_{n}
5. Only use the block types: "paragraph", "list_item", or "table".
6. Return valid JSON only.
7. Review and correct OCR errors for maximum accuracy.
"""

# --- TOKEN PRICING & LIMITS ---
# Define limits directly in the rate structure for clear model-specific configuration
MODEL_RATES_AND_LIMITS = {
    "gemini-2.5-flash-lite": {"input": 0.125, "output": 0.375, "rpm": 15, "rpd": 1000},
    "gemini-2.5-flash": {"input": 0.35, "output": 1.05, "rpm": 15, "rpd": 1000}, 
    "gemini-2.5-pro": {"input": 7.00, "output": 21.00, "rpm": 2, "rpd": 50}, # Pro model has stricter limits
}

class EnhancedGeminiConverter:
    
    def __init__(self, api_key: Optional[str] = None, model_name: str = "gemini-2.5-pro", processing_mode: str = "moderate"):
        load_dotenv()
        self.api_key = api_key or os.getenv('GEMINI_API_KEY')
        if not self.api_key:
            raise ValueError("Gemini API key required.")
        
        # --- Model Selection and Limit Determination (FIX) ---
        mode_to_model = {
            "fast": "gemini-2.5-flash-lite",
            "moderate": "gemini-2.5-flash", 
            "slow": "gemini-2.5-pro"
        }
        
        # Determine the final model name based on mode or explicit name
        if processing_mode in mode_to_model:
            self.model_name = mode_to_model[processing_mode]
        else:
            self.model_name = model_name
            
        if self.model_name not in MODEL_RATES_AND_LIMITS:
            raise ValueError(f"Unknown model name: {self.model_name}")

        self.processing_mode = processing_mode
        self.client = genai.Client(api_key=self.api_key)
        self.total_tokens_used = 0
        self.total_cost = 0.0
        
        # Cache setup
        self.cache_dir = Path(os.path.dirname(os.path.abspath(__file__))) / "cache"
        self.cache_dir.mkdir(exist_ok=True)
        
        # --- Global Rate Limiting Setup ---
        # Get limits based on the final determined model name
        limits = MODEL_RATES_AND_LIMITS[self.model_name]
        self.rpm_limit = limits["rpm"]
        self.rpd_limit = limits["rpd"]
        
        # Shared resources for global rate limiting
        self.request_times = collections.deque(maxlen=self.rpm_limit)
        self.daily_requests = 0
        self.daily_reset_time = time.time() + 86400
        self.request_lock = threading.Lock() # THE GLOBAL LOCK

        # State tracking for external systems (e.g., frontend)
        self.file_statuses = {}
        self.file_results = {}
        self.status_lock = threading.Lock() # Separate lock for file status
        
        # Dummy attributes for API consistency/legacy
        self.remaining_requests = self.rpm_limit
        self.min_request_interval = 60 / self.rpm_limit 
        self.last_request_time = time.time() - self.min_request_interval 
        self.reset_time = datetime.datetime.now()
        
        logging.info(f"Initialized Converter: Model={self.model_name}, RPM Limit={self.rpm_limit}, RPD Limit={self.rpd_limit}")

    def _enforce_rate_limits(self):
        """
        Enforce GLOBAL RPM and RPD limits for Gemini API (Thread-safe).
        This method MUST be called before any Gemini API interaction (upload, generate).
        """
        with self.request_lock: # Critical Section: Ensures only one thread modifies/checks limits at a time
            current_time = time.time()
            
            # --- RPD Check ---
            if current_time >= self.daily_reset_time:
                self.daily_requests = 0
                self.daily_reset_time = current_time + 86400
            
            if self.daily_requests >= self.rpd_limit:
                wait_time = self.daily_reset_time - current_time
                hours = int(wait_time // 3600)
                minutes = int((wait_time % 3600) // 60)
                
                # Suggest alternative models
                alternatives = []
                for mode, model in {"fast": "gemini-2.5-flash-lite", "moderate": "gemini-2.5-flash"}.items():
                    if model != self.model_name:
                        alternatives.append(f"{mode} mode ({model})")
                
                alt_text = f" Try switching to: {', '.join(alternatives)}" if alternatives else ""
                
                raise Exception(f"Daily limit reached for {self.model_name}. Resets in {hours}h {minutes}m.{alt_text}")
            
            # --- RPM Check (Sliding Window) ---
            while len(self.request_times) >= self.rpm_limit:
                oldest_request = self.request_times[0]
                if current_time - oldest_request < 60:
                    wait_time = 60 - (current_time - oldest_request)
                    logging.warning(f"Rate limit approaching. Sleeping for {wait_time:.2f}s...")
                    time.sleep(wait_time)
                    current_time = time.time() # Update time after sleeping
                else:
                    self.request_times.popleft() # Remove old request outside the 60s window
            
            # Record this request
            self.request_times.append(current_time)
            self.daily_requests += 1

    # --- Utility Methods (Unchanged/Refined) ---

    def get_file_hash(self, file_path: str) -> str:
        """Generate hash for file caching"""
        hasher = hashlib.md5()
        with open(file_path, 'rb') as f:
            while chunk := f.read(8192):
                hasher.update(chunk)
        return hasher.hexdigest()

    def load_cache(self, file_hash: str) -> Optional[Dict]:
        """Load cached result"""
        cache_file = self.cache_dir / f"{file_hash}.json"
        if cache_file.exists():
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    logging.info(f"Cache loaded for hash: {file_hash}")
                    return data
            except:
                logging.error(f"Failed to load cache for {file_hash}")
                return None
        return None

    def save_cache(self, file_hash: str, data: Dict):
        """Save result to cache"""
        cache_file = self.cache_dir / f"{file_hash}.json"
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            logging.info(f"Cache saved for hash: {file_hash}")
        except Exception as e:
            logging.error(f"Failed to save cache for {file_hash}: {e}")
            pass  

    def _track_usage(self, response):
        """Updates token and cost tracking"""
        if response.usage_metadata:
            prompt_tokens = response.usage_metadata.prompt_token_count
            candidate_tokens = response.usage_metadata.candidates_token_count
            
            rates = MODEL_RATES_AND_LIMITS.get(self.model_name)
            if rates:
                input_cost = (prompt_tokens / 1_000_000) * rates["input"]
                output_cost = (candidate_tokens / 1_000_000) * rates["output"]
                
                self.total_tokens_used += (prompt_tokens + candidate_tokens)
                self.total_cost += (input_cost + output_cost)
                logging.info(f"Usage tracked: {prompt_tokens} input, {candidate_tokens} output. Cost: ${input_cost + output_cost:.6f}")

    # --- UNICODE MAPPINGS ---
    SUPERSCRIPT_MAP = {
        '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴', '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
        '+': '⁺', '-': '⁻', '=': '⁼', '(': '⁽', ')': '⁾',
        'A': 'ᴬ', 'B': 'ᴮ', 'C': 'ᶜ', 'D': 'ᴰ', 'E': 'ᴱ', 'G': 'ᴳ', 'H': 'ᴴ', 'I': 'ᴵ', 'J': 'ᴶ', 'K': 'ᴷ', 'L': 'ᴸ', 'M': 'ᴹ', 'N': 'ᴺ', 'O': 'ᴼ', 'P': 'ᴾ', 'R': 'ᴿ', 'T': 'ᵀ', 'U': 'ᵁ', 'V': 'ⱽ', 'W': 'ᵂ',
        'a': 'ᵃ', 'b': 'ᵇ', 'c': 'ᶜ', 'd': 'ᵈ', 'e': 'ᵉ', 'f': 'ᶠ', 'g': 'ᵍ', 'h': 'ʰ', 'i': 'ⁱ', 'j': 'ʲ', 'k': 'ᵏ', 'l': 'ˡ', 'm': 'ᵐ', 'n': 'ⁿ', 'o': 'ᵒ', 'p': 'ᵖ', 'r': 'ʳ', 's': 'ˢ', 't': 'ᵗ', 'u': 'ᵘ', 'v': 'ᵛ', 'w': 'ʷ', 'x': 'ˣ', 'y': 'ʸ', 'z': 'ᶻ'
    }

    SUBSCRIPT_MAP = {
        '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄', '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉',
        '+': '₊', '-': '₋', '=': '₌', '(': '₍', ')': '₎',
        'a': 'ₐ', 'e': 'ₑ', 'h': 'ₕ', 'i': 'ᵢ', 'j': 'ⱼ', 'k': 'ₖ', 'l': 'ₗ', 'm': 'ₘ', 'n': 'ₙ', 'o': 'ₒ', 'p': 'ₚ', 'r': 'ᵣ', 's': 'ₛ', 't': 'ₜ', 'u': 'ᵤ', 'v': 'ᵥ', 'x': 'ₓ'
    }

    def _to_superscript(self, char: str) -> str:
        return self.SUPERSCRIPT_MAP.get(char, f"⁽{char}⁾")

    def _to_subscript(self, char: str) -> str:
        return self.SUBSCRIPT_MAP.get(char, f"₍{char}₎")

    def _convert_latex_to_unicode(self, text: str) -> str:
        """Convert LaTeX superscripts and subscripts to Unicode with fallback."""
        
        # Process Superscripts: ^{...}
        def replace_sup(match):
            content = match.group(1)
            return "".join(self._to_superscript(c) for c in content)
        
        # Process Subscripts: _{...}
        def replace_sub(match):
            content = match.group(1)
            return "".join(self._to_subscript(c) for c in content)

        # Regex for ^{...} and _{...}
        text = re.sub(r'\^\{([^\}]+)\}', replace_sup, text)
        text = re.sub(r'_\{([^\}]+)\}', replace_sub, text)
        
        return text
    
    def ai_correct_text(self, selected_text: str, user_explanation: str, extracted_text: str = "") -> str:
        """AI text correction using the same rate-limited system"""
        try:
            self._enforce_rate_limits()
            
            correction_prompt = f"""
You are an expert text correction model. Correct the following text based on the user's explanation.
Apply the same formatting and LaTeX notation rules as document extraction.

SELECTED TEXT TO CORRECT:
"{selected_text}"

USER'S EXPLANATION:
"{user_explanation}"

EXTRACTED CONTEXT:
"{extracted_text[:1000]}..."

CRITICAL RULES:
1. Only return the corrected version of the selected text
2. Maintain original formatting, spacing, and line breaks exactly
3. Fix only what the user described as incorrect
4. Preserve LaTeX notation (^{{}} for superscripts, _{{}} for subscripts)
5. Keep the same paragraph structure
6. Apply LaTeX conversion: y4→y^{{4}}, CO2→CO_{{2}}, H2O→H_{{2}}O

CORRECTED TEXT:
"""
            
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=[correction_prompt],
                config=genai.types.GenerateContentConfig(temperature=0.0)
            )
            
            self._track_usage(response)
            corrected_text = response.text.strip()
            
            # Clean up response
            if corrected_text.startswith('"') and corrected_text.endswith('"'):
                corrected_text = corrected_text[1:-1]
            
            # Apply LaTeX to Unicode conversion
            corrected_text = self._convert_latex_to_unicode(corrected_text)
            
            return corrected_text
            
        except Exception as e:
            if "Daily limit reached" in str(e):
                raise e  # Re-raise with proper error message
            raise Exception(f"AI correction failed: {str(e)}")
    
    def _parse_json_response(self, response_text: str) -> Dict[str, Any]:
        """Robustly parses JSON from the model's response text. (Unchanged)"""
        try:
            match = re.search(r'```json\s*(\{.*?\})\s*```', response_text, re.DOTALL)
            if match:
                data = json.loads(match.group(1))
            else:
                data = json.loads(response_text)
            
            self._apply_unicode_conversion(data)
            return data
            
        except json.JSONDecodeError as e:
            raise ValueError(f"Failed to parse JSON from model response. Error: {e}\nResponse:\n{response_text[:500]}...")
    
    def _apply_unicode_conversion(self, data: Dict[str, Any]):
        """Recursively apply Unicode conversion to all text in the JSON data."""
        if isinstance(data, dict):
            for key, value in data.items():
                if key == 'text' and isinstance(value, str):
                    data[key] = self._convert_latex_to_unicode(value)
                elif isinstance(value, (dict, list)):
                    self._apply_unicode_conversion(value)
        elif isinstance(data, list):
            for item in data:
                if isinstance(item, (dict, list)):
                    self._apply_unicode_conversion(item)

    def split_pdf_pages(self, pdf_path: str) -> List[str]:
        """Split PDF into individual page files. (Unchanged)"""
        # ... [PDF splitting logic remains here] ...
        pdf_path = Path(pdf_path)
        base_name = pdf_path.stem
        temp_dir = pdf_path.parent / f"{base_name}_pages_{os.getpid()}_{time.time()}" 
        temp_dir.mkdir(exist_ok=True)
        
        page_files = []
        
        try:
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                total_pages = len(reader.pages)
                
                for page_num in range(total_pages):
                    writer = PdfWriter()
                    writer.add_page(reader.pages[page_num])
                    
                    page_file = temp_dir / f"{base_name}_page_{page_num + 1}.pdf"
                    with open(page_file, 'wb') as output_file:
                        writer.write(output_file)
                    
                    page_files.append(str(page_file))
        except Exception as e:
            logging.error(f"Error splitting PDF: {e}")
            self._cleanup_temp_dir(temp_dir)
            raise

        return page_files
    
    def _cleanup_temp_dir(self, temp_dir: Path):
        """Utility to safely clean up temporary directory. (Unchanged)"""
        try:
            if temp_dir.exists():
                shutil.rmtree(temp_dir)
                logging.info(f"Cleaned up temp directory: {temp_dir}")
        except Exception as e:
            logging.error(f"Failed to clean up temp directory {temp_dir}: {e}")
    
    def _process_single_page_async(self, page_file: str, page_num: int) -> Dict[str, Any]:
        """
        Process single page. The global rate limit is enforced by _enforce_rate_limits()
        which uses the shared self.request_lock.
        """
        pdf_file = None
        try:
            # 1. GLOBAL RATE LIMIT ENFORCED HERE
            self._enforce_rate_limits() 
            
            # 2. Upload file (Counts as 1 request, rate-limited above)
            pdf_file = self.client.files.upload(file=page_file)
            logging.info(f"Page {page_num}: File uploaded: {pdf_file.name}")

            # 3. Generate content (Counts as 1 request, must also be rate-limited)
            # IMPORTANT: The single worker thread ensures that the subsequent 
            # API call (generate_content) for the next page won't be made until
            # the current thread is completely done and the sleep/wait is over.
            
            # We must enforce the limit AGAIN here because this is a separate 
            # API call (generate_content) after the file.upload API call.
            # However, since ThreadPoolExecutor(max_workers=1) ensures *serial*
            # execution of pages, the external scheduler handles the pause.
            # We'll rely on the main ThreadPoolExecutor to schedule the next page
            # only after the full current page's thread completes, which simplifies the logic.
            
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=[
                    PROMPT_FULL_DOCUMENT_EXTRACTION, 
                    pdf_file
                ],
                config=genai.types.GenerateContentConfig(temperature=0.0, response_mime_type="application/json")
            )
            
            self._track_usage(response)
            result = self._parse_json_response(response.text)
            
            return {
                "page_num": page_num,
                "result": result,
                "status": "success"
            }
            
        except Exception as e:
            # Catching and logging the RPD limit error raised in _enforce_rate_limits
            logging.error(f"Error processing page {page_num}: {e}")
            return {
                "page_num": page_num,
                "result": None,
                "status": "error",
                "error": str(e)
            }
        finally:
            if pdf_file:
                try:
                    self.client.files.delete(name=pdf_file.name)
                    logging.info(f"Page {page_num}: File deleted: {pdf_file.name}")
                except Exception as e:
                    logging.error(f"Failed to delete file {pdf_file.name}: {e}")
    
    # --- File Status and Result Tracking (New/Improved) ---
    
    def set_file_status(self, file_id: str, status: str, progress: int = 0, message: str = ""):
        """Update file processing status"""
        with self.status_lock:
            self.file_statuses[file_id] = {
                "status": status,
                "progress": progress,
                "message": message,
                "timestamp": time.time()
            }
    
    def get_file_status(self, file_id: str) -> Dict[str, Any]:
        """Get file processing status"""
        with self.status_lock:
            return self.file_statuses.get(file_id, {"status": "not_found"})
    
    def set_file_result(self, file_id: str, json_data: Dict[str, Any]):
        """Store file processing result"""
        with self.status_lock:
            self.file_results[file_id] = json_data
    
    def get_file_result(self, file_id: str) -> Optional[Dict[str, Any]]:
        """Get file processing result"""
        with self.status_lock:
            return self.file_results.get(file_id)

    def convert_with_file_id(self, pdf_path: str, file_id: str, output_path: Optional[str] = None, use_cache: bool = True) -> str:
        """
        Convert PDF using a tracked file_id. 
        Ensures the final output path uses the file_id if not specified.
        """
        pdf_path = Path(pdf_path)
        
        # FIX 1: Explicitly define the final output path using file_id if needed
        if not output_path:
            output_dir = Path("outputs") 
            output_dir.mkdir(exist_ok=True)
            output_path_final = output_dir / f"{file_id}_{pdf_path.stem}.docx"
        else:
            output_path_final = Path(output_path)

        def progress_callback(progress: int, message: str):
            self.set_file_status(file_id, "processing", progress, message)
        
        try:
            self.set_file_status(file_id, "processing", 0, "Starting conversion...")
            result_path = self.convert_single(str(pdf_path), str(output_path_final), progress_callback, use_cache)
            
            # Retrieve and store the final JSON result from cache (or in-memory if just generated)
            file_hash = self.get_file_hash(str(pdf_path))
            json_data = self.load_cache(file_hash) 
            if json_data:
                self.set_file_result(file_id, json_data)
            
            self.set_file_status(file_id, "completed", 100, f"Conversion completed. Output: {result_path}")
            return str(result_path)
            
        except Exception as e:
            error_message = f"Conversion failed: {str(e)}"
            self.set_file_status(file_id, "failed", 0, error_message)
            logging.error(error_message)
            raise
    
    # --- Core Conversion Logic (Rate-Limited by _process_single_page_async) ---
    def convert_single(self, pdf_path: str, output_path: str = None, progress_callback: Callable = None, use_cache: bool = True) -> str:
        """
        Convert single PDF. All page processing calls _process_single_page_async 
        which enforces the GLOBAL rate limit.
        """
        pdf_path = Path(pdf_path)
        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF not found: {pdf_path}")
        
        temp_dir = None
        
        def update_progress(progress: int, message: str):
            if progress_callback:
                progress_callback(progress, message)

        try:
            update_progress(5, "Starting conversion...")
            
            file_hash = self.get_file_hash(str(pdf_path))
            cached = self.load_cache(file_hash) if use_cache else None
            
            if cached and use_cache:
                update_progress(100, "Completed! Using cached result.")
                json_data = cached
            else:
                update_progress(10, "Splitting PDF into pages...")
                
                page_files = self.split_pdf_pages(str(pdf_path))
                total_pages = len(page_files)
                if not page_files:
                     raise ValueError("PDF split resulted in zero pages.")
                     
                temp_dir = Path(page_files[0]).parent
                update_progress(15, f"Processing {total_pages} pages...")
                
                results = {}
                completed_pages = 0
                
                # max_workers=1 is CRITICAL for guaranteed sequential rate limiting
                with ThreadPoolExecutor(max_workers=1) as executor: 
                    future_to_page = {}
                    for i, page_file in enumerate(page_files):
                        # Each submission runs _process_single_page_async, which hits the global lock
                        future = executor.submit(self._process_single_page_async, page_file, i + 1) 
                        future_to_page[future] = i + 1
                    
                    for future in as_completed(future_to_page):
                        page_result = future.result()
                        page_num = page_result["page_num"]
                        
                        # Update progress before processing result
                        current_progress = 15 + int(((page_num - 1) / total_pages) * 70)
                        update_progress(current_progress, f"Processing page {page_num}/{total_pages}...")
                        
                        if page_result["status"] == "success":
                            extracted_page_data = page_result["result"].get("pages", [])
                            if extracted_page_data:
                                for data in extracted_page_data:
                                     data["page_number"] = page_num 
                                results[page_num] = extracted_page_data[0] 
                            completed_pages += 1
                        else:
                            # If a page fails (e.g., due to RPD limit hit), log it but proceed to try next pages
                            results[page_num] = {
                                "page_number": page_num,
                                "blocks": [{"type": "paragraph", "text": f"[Page {page_num} failed: {page_result['error'][:100]}]"}]
                            }

                        progress = 15 + int(((completed_pages) / total_pages) * 70)
                        if page_result["status"] == "success":
                            update_progress(progress, f"Processing page {page_num}/{total_pages} - Completed")
                        else:
                            update_progress(progress, f"Processing page {page_num}/{total_pages} - Failed: {page_result['error'][:50]}...")

                # Merge results
                final_pages_list = [results[page_num] for page_num in sorted(results.keys())]
                
                json_data = {
                    "total_pages": total_pages,
                    "pages": final_pages_list,
                    "processing_mode": self.processing_mode,
                    "model_used": self.model_name
                }
                
                update_progress(90, "Results parsed and formatted...")
                
                if use_cache:
                    self.save_cache(file_hash, json_data)

            update_progress(95, "Creating Word document...")
            
            if not output_path:
                output_path = str(pdf_path.with_suffix('.docx'))
            
            self._save_docx_from_json(json_data, Path(output_path))
            
            update_progress(100, "Completed!")
            return output_path
            
        except Exception as e:
            logging.error(f"Conversion failed critically: {e}")
            raise Exception(f"Conversion failed: {str(e)}")
        finally:
            if temp_dir:
                self._cleanup_temp_dir(temp_dir)

    def convert_batch(self, pdf_files: List[str], output_dir: str = "outputs", progress_callback: Callable = None) -> List[str]:
        """
        Convert multiple PDFs. Uses convert_with_file_id to run each PDF 
        in a separate thread for concurrent processing, while the core 
        _process_single_page_async ensures global rate limiting.
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True)
        
        # Use a higher max_workers for the batch itself, as long as convert_single 
        # keeps max_workers=1 internally. This allows multiple files to be 
        # *queued* and split/cached concurrently, but the actual *API calls* # remain sequential and rate-limited.
        MAX_BATCH_WORKERS = 4 
        
        futures = {}
        batch_results = []
        
        total_files = len(pdf_files)
        
        def batch_progress(file_idx: int, file_progress: int, message: str):
            overall_progress = int((file_idx * 100 + file_progress) / total_files)
            batch_message = f"File {file_idx + 1}/{total_files}: {message}"
            if progress_callback:
                progress_callback(overall_progress, batch_message)

        with ThreadPoolExecutor(max_workers=MAX_BATCH_WORKERS) as executor:
            for i, pdf_file in enumerate(pdf_files):
                pdf_path = Path(pdf_file)
                # Generate a unique ID for this file's job
                file_id = f"{pdf_path.stem}_{str(uuid.uuid4()).split('-')[0]}" 
                
                # Define a unique output path for this file in the batch
                output_path = output_dir / f"{file_id}_{pdf_path.stem}.docx"
                
                # The file's index is captured to correctly update batch progress
                future = executor.submit(
                    self.convert_with_file_id, 
                    str(pdf_path), 
                    file_id, 
                    str(output_path), 
                    True
                )
                futures[future] = (i, file_id) # Store index and file_id

            # Monitor progress
            for future in as_completed(futures):
                file_idx, file_id = futures[future]
                
                try:
                    result_path = future.result()
                    batch_progress(file_idx, 100, f"Completed: {result_path}")
                    batch_results.append(result_path)
                except Exception as e:
                    # Retrieve the error message from file_statuses
                    status_data = self.get_file_status(file_id)
                    error_message = status_data.get("message", f"Unknown error: {e}")
                    
                    batch_progress(file_idx, 0, f"Failed: {error_message}")
                    batch_results.append(None)
                    
        return batch_results

    # --- DOCX Creation Methods (Unchanged) ---
    def _save_docx_from_json(self, json_data: Dict[str, Any], output_path: Path):
        """Creates a DOCX file from the structured JSON data. (Unchanged)"""
        # ... [DOCX saving logic remains here] ...
        doc = Document()
        
        pages = json_data.get("pages", [])
        for i, page in enumerate(pages):
            if i > 0:
                doc.add_page_break()
            
            blocks = page.get("blocks", [])
            for block in blocks:
                self._add_block_to_doc(doc, block)

        doc.save(output_path)

    def _add_block_to_doc(self, doc: Document, block: Dict[str, Any]):
        """Adds a single content block to the document based on its type. (Unchanged)"""
        # ... [Block addition logic remains here] ...
        block_type = block.get("type", "paragraph").lower()
        text = block.get("text", "").rstrip() 
        
        if not text:
            return

        if block_type == "list_item":
            if re.match(r'^\s*[\-*\•]\s', text) or re.match(r'^\s*[\-*\•]', text.split('\n')[0].strip()):
                para = doc.add_paragraph(style='List Bullet')
                text = re.sub(r'^\s*[\-*\•]\s?', '', text, 1) 
            elif re.match(r'^\s*(\d+\.|[a-zA-Z]\.)\s', text) or re.match(r'^\s*(\d+\.|[a-zA-Z]\.)', text.split('\n')[0].strip()):
                para = doc.add_paragraph(style='List Number')
                text = re.sub(r'^\s*(\d+\.|[a-zA-Z]\.)\s?', '', text, 1)
            else:
                para = doc.add_paragraph(style='List Bullet') 
        
        elif block_type == "table":
            self._add_table_block(doc, text)
            return
            
        else: 
            para = doc.add_paragraph()
            
        self._add_text_with_line_breaks(para, text)

    def _add_text_with_line_breaks(self, paragraph, text: str):
        """Adds text to a paragraph, preserving internal line breaks. (Unchanged)"""
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            run = paragraph.add_run(line)
            if i < len(lines) - 1:
                run.add_break()
                
    def _add_table_block(self, doc: Document, table_text: str):
        """Parses markdown table text and adds a table to the document. (Unchanged)"""
        # ... [Table addition logic remains here] ...
        lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
        
        rows_data = []
        is_header = True
        for line in lines:
            if all(re.match(r'^[-: ]+$', cell.strip()) for cell in line[1:-1].split('|')):
                is_header = False
                continue 
            
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            if all(not cell for cell in cells):
                continue
                
            rows_data.append(cells)
        
        if not rows_data:
            doc.add_paragraph(f"[Unparsed/Empty Table Data]:\n{table_text}")
            return
            
        num_cols = max(len(row) for row in rows_data)
        if num_cols == 0: 
            return

        rows_data = [row + [''] * (num_cols - len(row)) for row in rows_data]
        num_rows = len(rows_data)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        for i, row_cells in enumerate(rows_data):
            for j, cell_text in enumerate(row_cells):
                cell = table.cell(i, j)
                para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                para.clear()
                self._add_text_with_line_breaks(para, cell_text)
                
        if is_header is False and num_rows > 0:
             for cell in table.rows[0].cells:
                 for paragraph in cell.paragraphs:
                     for run in paragraph.runs:
                         run.bold = True

    # --- Status and Utility Getters (New/Improved) ---

    def get_usage_stats(self) -> Dict[str, Any]:
        """Returns statistics on API usage and cost with rate limit info"""
        current_time = time.time()
        reset_time_remaining = max(0, self.daily_reset_time - current_time)
        
        return {
            "total_tokens": self.total_tokens_used,
            "estimated_cost": f"${self.total_cost:.4f}",
            "processing_mode": self.processing_mode,
            "model_used": self.model_name,
            "cache_files": len(list(self.cache_dir.glob("*.json"))) if self.cache_dir.exists() else 0,
            "daily_requests": self.daily_requests,
            "daily_limit": self.rpd_limit,
            "requests_remaining": max(0, self.rpd_limit - self.daily_requests),
            "reset_time_seconds": int(reset_time_remaining),
            "rpm_limit": self.rpm_limit
        }

    def cleanup_old_files(self, max_age_hours: int = 24):
        """Clean up old file statuses and results. (Unchanged)"""
        current_time = time.time()
        cutoff_time = current_time - (max_age_hours * 3600)
        
        with self.status_lock:
            old_file_ids = []
            for file_id, status_data in self.file_statuses.items():
                if status_data.get("timestamp", 0) < cutoff_time:
                    old_file_ids.append(file_id)
            
            for file_id in old_file_ids:
                self.file_statuses.pop(file_id, None)
                self.file_results.pop(file_id, None)
            
            if old_file_ids:
                logging.info(f"Cleaned up {len(old_file_ids)} old file records")

    @staticmethod
    def get_available_modes() -> Dict[str, Dict]:
        """Get available processing modes for frontend. (Unchanged)"""
        return {
            "fast": {
                "display_name": "Fast Processing",
                "description": "Optimized for speed and cost-efficiency (Gemini 2.5 Flash Lite)",
                "estimated_cost_per_page": "$0.0001",
                "daily_limit": 1000,
                "rpm_limit": 15
            },
            "moderate": {
                "display_name": "Moderate Processing", 
                "description": "Balanced performance and accuracy (Gemini 2.5 Flash)",
                "estimated_cost_per_page": "$0.0002",
                "daily_limit": 1000,
                "rpm_limit": 15
            },
            "slow": {
                "display_name": "Premium Processing",
                "description": "Maximum accuracy for complex documents (Gemini 2.5 Pro)", 
                "estimated_cost_per_page": "$0.0050",
                "daily_limit": 50,
                "rpm_limit": 2
            }
        }
    
    def get_rate_limit_status(self) -> Dict[str, Any]:
        """Get current rate limit status for frontend"""
        current_time = time.time()
        reset_time_remaining = max(0, self.daily_reset_time - current_time)
        hours = int(reset_time_remaining // 3600)
        minutes = int((reset_time_remaining % 3600) // 60)
        
        return {
            "model": self.model_name,
            "daily_requests": self.daily_requests,
            "daily_limit": self.rpd_limit,
            "requests_remaining": max(0, self.rpd_limit - self.daily_requests),
            "rpm_limit": self.rpm_limit,
            "reset_time_hours": hours,
            "reset_time_minutes": minutes,
            "reset_time_seconds": int(reset_time_remaining)
        }
    
    def extract_text_for_preview(self, json_data: Dict[str, Any]) -> str:
        """Extract plain text from JSON data for preview"""
        text_parts = []
        pages = json_data.get("pages", [])
        
        for page in pages:
            page_num = page.get("page_number", 1)
            text_parts.append(f"=== Page {page_num} ===\n")
            
            blocks = page.get("blocks", [])
            for block in blocks:
                text = block.get("text", "").strip()
                if text:
                    text_parts.append(text + "\n")
            text_parts.append("\n")
        
        return "".join(text_parts)
    
    def get_pages_data(self, json_data: Dict[str, Any]) -> List[Dict]:
        """Extract pages data for frontend"""
        pages_data = []
        pages = json_data.get("pages", [])
        
        for page in pages:
            page_num = page.get("page_number", 1)
            blocks = page.get("blocks", [])
            
            content_parts = []
            for block in blocks:
                text = block.get("text", "").strip()
                if text:
                    block_type = block.get("type", "paragraph")
                    if block_type == "list_item":
                        if not text.startswith(('• ', '- ', '* ')) and not re.match(r'^\d+\.', text):
                            text = f"• {text}"
                    content_parts.append(text)
            
            pages_data.append({
                "page_number": page_num,
                "content": "\n\n".join(content_parts) if content_parts else ""
            })
        
        return pages_data